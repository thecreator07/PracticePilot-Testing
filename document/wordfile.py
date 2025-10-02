from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import List
from pydantic_class import LLMResponseModel

# horizontal line function
def add_hr(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bottom = OxmlElement("w:pBdr")
        border = OxmlElement("w:bottom")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "auto")
        bottom.append(border)
        pPr.append(bottom)



def create_patient_docx(llm_response: LLMResponseModel, output_path: str):
    doc = Document()
    
    # Title
    heading=doc.add_heading("Patient Medical Report", level=0)
    heading.alignment = 1  # Center alignment
    heading.runs[0].font.size = 240000 
    heading.runs[0].font.bold = True    
    
    add_hr(heading) 
    def add_section(title: str, content: List[str], bullet: bool = False, emoji: str = None):
        if(title!="none"):            
            header = doc.add_heading(title, level=2)
            header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if content:
            for item in content:
                if bullet:
                    if emoji:
                        doc.add_paragraph(f"{emoji} {item}", style="Normal")
                    else:
                        doc.add_paragraph(f"• {item}", style="Normal")
                else:
                    doc.add_paragraph(item, style="Normal")
        else:
            doc.add_paragraph("(No data)", style="Italic")

        if title!="none"and content.__len__()>1:
            doc.add_paragraph("") # Add a blank line after each section
            
            
            
  
    soapheader=doc.add_heading("SOAP Notes", level=1)
    add_hr(soapheader) 
    add_section("Subjective", [llm_response.SOAP.Subjective] if llm_response.SOAP.Subjective else [])    
    add_section("Objective", [llm_response.SOAP.Objective] if llm_response.SOAP.Objective else [])
    add_section("Assessment", [llm_response.SOAP.Assessment] if llm_response.SOAP.Assessment else [])
    add_section("Plan", llm_response.SOAP.Plan, bullet=True,emoji="➡️")
   
    clinicianheader=doc.add_heading("Clinician Summary", level=1)
    add_hr(clinicianheader) 
    add_section("Patient Details", [llm_response.Clinician_Summary.Patient_Details] if llm_response.Clinician_Summary.Patient_Details else [])
    add_section("Chief Complaint", [llm_response.Clinician_Summary.Chief_Complaint] if llm_response.Clinician_Summary.Chief_Complaint else [])
    add_section("History", [llm_response.Clinician_Summary.History] if llm_response.Clinician_Summary.History else [])
    add_section("Assessment", [llm_response.Clinician_Summary.Assessment] if llm_response.Clinician_Summary.Assessment else [])
    add_section("Plan", llm_response.Clinician_Summary.Plan, bullet=True,emoji="➡️")
    
    patientheader=doc.add_heading("Patient Summary", level=1)
    add_hr(patientheader)
    
    add_section("Clinical Actions", llm_response.Patient_Summary.What_We_Did, bullet=True,emoji="✅")
    add_section("Expected Course", llm_response.Patient_Summary.What_To_Expect, bullet=True,emoji="🔮")
    doc.add_heading("Self-Care Guidelines", level=2)
    if llm_response.Patient_Summary.Dos_and_Donts.Dos or llm_response.Patient_Summary.Dos_and_Donts.Donts:
        if llm_response.Patient_Summary.Dos_and_Donts.Dos:
           add_section("none", llm_response.Patient_Summary.Dos_and_Donts.Dos, bullet=True,emoji="✅")
        if llm_response.Patient_Summary.Dos_and_Donts.Donts:
           add_section("none", llm_response.Patient_Summary.Dos_and_Donts.Donts, bullet=True,emoji="❌")
    doc.add_paragraph("") 
    add_section("Emergency Contact Guidelines", llm_response.Patient_Summary.When_To_Call, bullet=True,emoji="📞")
    add_section("Continuity of Care", llm_response.Patient_Summary.Next_Steps, bullet=True,emoji="⏭️")

    # Save document
    doc.save(output_path)
    print(f"Document saved to {output_path}")


